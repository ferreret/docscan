#!/usr/bin/env python3.14
"""Genera el manual de usuario de DocScan Studio en formato .docx.

Uso:
    python3.14 docs/generate_manual.py
    Salida: docs/manual_docscan_studio.docx

Requiere: pip install python-docx pillow
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    from PIL import Image, ImageDraw
except ImportError as exc:
    sys.exit(f"Dependencia faltante: {exc}\nInstala con: pip install python-docx pillow")

OUTPUT = Path(__file__).parent / "manual_docscan_studio.docx"

# Colores corporativos
_BLUE_DARK = RGBColor(0x1F, 0x39, 0x7D)
_BLUE_MID = RGBColor(0x2E, 0x74, 0xB5)
_BLUE_LIGHT = RGBColor(0x44, 0x72, 0xC4)
_GRAY = RGBColor(0x59, 0x59, 0x59)

_VERSION = "0.1 RC"
_DATE = "Marzo 2026"

# ── Helpers ─────────────────────────────────────────────────────────


def _cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _page_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f" {field_code} "
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    run._r.extend([fc1, it, fc2])


def _placeholder(label: str, w: int = 800, h: int = 350) -> io.BytesIO:
    """Genera un rectángulo gris con texto centrado como PNG en memoria."""
    img = Image.new("RGB", (w, h), (220, 220, 220))
    draw = ImageDraw.Draw(img)
    draw.rectangle([2, 2, w - 3, h - 3], outline=(160, 160, 160), width=2)
    # Icono de cámara simple
    cx, cy = w // 2, h // 2 - 15
    draw.rectangle([cx - 30, cy - 20, cx + 30, cy + 20], outline=(120, 120, 120), width=2)
    draw.ellipse([cx - 12, cy - 12, cx + 12, cy + 12], outline=(120, 120, 120), width=2)
    # Texto
    bbox = draw.textbbox((0, 0), label)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((w - tw) // 2, cy + 35), label, fill=(80, 80, 80))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def _figure(doc: Document, label: str, caption: str, w_cm: float = 14) -> None:
    """Añade imagen placeholder con pie de figura."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(_placeholder(label), width=Cm(w_cm))
    doc.add_paragraph(caption, style="Caption")


def _table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    caption: str = "",
) -> None:
    """Tabla con cabecera coloreada y filas alternas."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        _cell_bg(cell, "1F397D")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, rd in enumerate(rows):
        rc = t.add_row().cells
        bg = "DEEAF1" if idx % 2 == 0 else "FFFFFF"
        for i, v in enumerate(rd):
            rc[i].text = v
            _cell_bg(rc[i], bg)
    if caption:
        doc.add_paragraph(caption, style="Caption")


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


# ── Documento ───────────────────────────────────────────────────────


def build() -> None:
    doc = Document()

    # ── Estilos ──
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = _BLUE_DARK
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(10)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = _BLUE_MID
    h2.paragraph_format.space_before = Pt(14)

    h3 = doc.styles["Heading 3"]
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = _BLUE_LIGHT

    cap = doc.styles["Caption"]
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = _GRAY
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(12)

    # ── Página ──
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    sec.different_first_page_header_footer = True

    # Cabecera
    hdr = sec.header.paragraphs[0]
    hdr.clear()
    r = hdr.add_run("DocScan Studio — Manual de Usuario")
    r.font.size = Pt(9)
    r.font.color.rgb = _GRAY
    hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Pie de página
    ftr = sec.footer.paragraphs[0]
    ftr.clear()
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ftr.add_run("Página ").font.size = Pt(9)
    _page_field(ftr, "PAGE")
    ftr.add_run(" de ").font.size = Pt(9)
    _page_field(ftr, "NUMPAGES")

    # ════════════════════════════════════════════════════════════════
    # PORTADA
    # ════════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DocScan Studio")
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = _BLUE_DARK

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Manual de Usuario")
    r2.font.size = Pt(18)
    r2.font.color.rgb = _BLUE_MID

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Versión {_VERSION} — {_DATE}").font.color.rgb = _GRAY

    doc.add_paragraph()
    _figure(doc, "[Logo DocScan Studio]", "", w_cm=8)

    for _ in range(3):
        doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("© 2026 Tecnomedia")
    r4.font.size = Pt(10)
    r4.font.color.rgb = _GRAY

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # TABLA DE CONTENIDOS
    # ════════════════════════════════════════════════════════════════
    toc_h = doc.add_paragraph("Tabla de contenidos", style="Heading 1")
    toc_h.paragraph_format.page_break_before = False

    toc_p = doc.add_paragraph()
    r = toc_p.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = 'TOC \\o "1-3" \\h \\z \\u'
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r._r.extend([fc1, it, fc2])

    doc.add_paragraph(
        "(Pulsar Ctrl+A → F9 en Word para actualizar la tabla de contenidos)",
    ).italic = True

    # ════════════════════════════════════════════════════════════════
    # CAPÍTULO 1 — INTRODUCCIÓN
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph("1. Introducción", style="Heading 1")

    doc.add_paragraph(
        "DocScan Studio es una plataforma de escritorio para la captura por lotes, "
        "procesamiento e indexación de documentos. Diseñada para entornos de "
        "digitalización documental con alto volumen, combina un pipeline configurable "
        "con inteligencia artificial generativa para automatizar tareas complejas."
    )

    doc.add_paragraph("1.1 ¿Para quién es DocScan Studio?", style="Heading 2")
    _bullets(doc, [
        "Departamentos de administración y archivo que digitalizan facturas, albaranes y contratos",
        "Empresas de gestión documental que procesan documentos para terceros",
        "Cualquier organización que necesite capturar, clasificar y exportar documentos escaneados",
    ])

    doc.add_paragraph("1.2 Características principales", style="Heading 2")
    _bullets(doc, [
        "Multi-aplicación: cada perfil de procesamiento es completamente independiente",
        "Pipeline composable: combinación libre de pasos de imagen, barcode, OCR y scripts",
        "Doble motor de barcodes: pyzbar (rápido) y zxing-cpp (14 simbologías)",
        "Triple motor OCR: RapidOCR (offline), EasyOCR (alta precisión) y Tesseract (ligero)",
        "Integración IA: Claude (Anthropic) y GPT-4o (OpenAI) para clasificación y extracción",
        "Scripting Python integrado con acceso completo al contexto de la aplicación",
        "AI Mode: asistente conversacional para crear y configurar aplicaciones",
        "Internacionalización: Español, English, Català",
        "Temas claro y oscuro con atajos de teclado configurables",
        "Multiplataforma: Linux (SANE) y Windows (TWAIN/WIA)",
    ])

    doc.add_paragraph("1.3 Flujo de trabajo típico", style="Heading 2")
    _numbered(doc, [
        "Crear una aplicación en el Launcher con la configuración deseada",
        "Abrir la aplicación para iniciar un nuevo lote",
        "Escanear documentos o importar ficheros (TIFF, JPEG, PNG, PDF, BMP)",
        "El pipeline procesa automáticamente cada página (corrección, barcodes, OCR, scripts)",
        "Revisar los resultados en el Workbench: barcodes detectados, texto OCR, campos",
        "Transferir el lote al sistema de destino (carpeta, PDF, CSV, transferencia avanzada)",
    ])

    _figure(
        doc,
        "[SCREENSHOT: Flujo de trabajo — desde escaneo hasta transferencia]",
        "Figura 1. Flujo de trabajo completo de DocScan Studio.",
    )

    # ════════════════════════════════════════════════════════════════
    # CAPÍTULO 2 — INSTALACIÓN
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph("2. Instalación", style="Heading 1")

    doc.add_paragraph("2.1 Requisitos del sistema", style="Heading 2")
    _table(doc, ["Componente", "Requisito", "Notas"], [
        ["Sistema operativo", "Linux / Windows 10+", "macOS no soportado oficialmente"],
        ["Python", "3.14 o superior", "En Linux: python3.14"],
        ["RAM", "4 GB mínimo", "8 GB recomendado para OCR + IA"],
        ["Disco", "500 MB + espacio para lotes", "Los modelos OCR ocupan ~10–500 MB"],
        ["Escáner", "Compatible SANE (Linux) o TWAIN/WIA (Windows)", "Opcional, se puede usar solo importación"],
    ], "Tabla 1. Requisitos del sistema.")

    doc.add_paragraph("2.2 Instalación paso a paso", style="Heading 2")
    _numbered(doc, [
        "Clonar el repositorio: git clone https://github.com/ferreret/docscan.git",
        "Crear entorno virtual: python3.14 -m venv .venv",
        "Activar el entorno: source .venv/bin/activate (Linux) o .venv\\Scripts\\activate (Windows)",
        "Instalar dependencias: pip install -r requirements.txt",
        "Inicializar la base de datos: alembic upgrade head",
        "Ejecutar: python3.14 main.py",
    ])

    _figure(
        doc,
        "[SCREENSHOT: Terminal con el proceso de instalación completado]",
        "Figura 2. Instalación completada en terminal.",
    )

    doc.add_paragraph("2.3 Primer arranque", style="Heading 2")
    doc.add_paragraph(
        "Al ejecutar DocScan Studio por primera vez, se muestra una pantalla de splash "
        "con el progreso de inicialización. La base de datos SQLite se crea automáticamente "
        "y el Launcher aparece vacío, listo para crear la primera aplicación."
    )

    _figure(
        doc,
        "[SCREENSHOT: Splash screen de DocScan Studio]",
        "Figura 3. Pantalla de splash durante el arranque.",
    )

    # ════════════════════════════════════════════════════════════════
    # CAPÍTULO 3 — LAUNCHER
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph("3. Launcher", style="Heading 1")

    doc.add_paragraph(
        "El Launcher es la ventana principal de DocScan Studio. Desde aquí se gestionan "
        "todas las aplicaciones (perfiles de procesamiento) y se accede al resto de funcionalidades."
    )

    _figure(
        doc,
        "[SCREENSHOT: Launcher con la sidebar colapsada y lista de aplicaciones]",
        "Figura 4. Launcher de DocScan Studio con sidebar colapsada.",
    )

    doc.add_paragraph("3.1 Sidebar", style="Heading 2")
    doc.add_paragraph(
        "La sidebar izquierda proporciona acceso rápido a todas las acciones. "
        "Se puede expandir haciendo clic en el icono de menú para ver los nombres "
        "de cada acción junto a los iconos."
    )

    _table(doc, ["Botón", "Acción", "Descripción"], [
        ["Nueva", "Crear aplicación", "Abre el diálogo para crear una nueva aplicación"],
        ["Abrir", "Abrir aplicación", "Abre el workbench para la aplicación seleccionada"],
        ["Configurar", "Editar configuración", "Abre el configurador con las 6 pestañas"],
        ["Clonar", "Duplicar aplicación", "Crea una copia exacta de la aplicación seleccionada"],
        ["Exportar", "Exportar como .docscan", "Guarda la configuración en un fichero JSON portable"],
        ["Importar", "Importar .docscan", "Carga una aplicación desde fichero exportado"],
        ["Eliminar", "Eliminar aplicación", "Elimina la aplicación y todos sus datos (irreversible)"],
        ["Actualizar", "Refrescar lista", "Recarga la lista de aplicaciones desde la base de datos"],
        ["Gestor de Lotes", "Histórico", "Abre el gestor de lotes para consultar el historial"],
        ["AI MODE", "Asistente IA", "Activa el panel conversacional con inteligencia artificial"],
        ["Acerca de", "Información", "Muestra la versión y los créditos de la aplicación"],
    ], "Tabla 2. Acciones de la sidebar.")

    _figure(
        doc,
        "[SCREENSHOT: Sidebar expandida mostrando todos los botones con texto]",
        "Figura 5. Sidebar expandida con todas las acciones visibles.",
    )

    doc.add_paragraph("3.2 Crear una nueva aplicación", style="Heading 2")
    doc.add_paragraph(
        "Al hacer clic en «Nueva», aparece un diálogo donde se introduce el nombre "
        "y la descripción de la aplicación. Una vez creada, aparece en la lista "
        "y se puede configurar mediante el botón «Configurar»."
    )

    _figure(
        doc,
        "[SCREENSHOT: Diálogo de nueva aplicación con nombre y descripción]",
        "Figura 6. Diálogo para crear una nueva aplicación.",
    )

    doc.add_paragraph("3.3 Exportar e importar aplicaciones", style="Heading 2")
    doc.add_paragraph(
        "Las aplicaciones se pueden exportar como ficheros .docscan (JSON) para "
        "compartirlas entre instalaciones o como copia de seguridad. El fichero "
        "incluye toda la configuración: pipeline, campos, eventos y transferencia."
    )

    doc.add_paragraph("3.4 Tema y idioma", style="Heading 2")
    doc.add_paragraph(
        "La barra superior del Launcher permite alternar entre tema claro y oscuro "
        "y cambiar el idioma de la interfaz (Español, English, Català). El cambio "
        "de idioma requiere reiniciar la aplicación."
    )

    _figure(
        doc,
        "[SCREENSHOT: Launcher en tema oscuro]",
        "Figura 7. Launcher con tema oscuro activado.",
    )

    # ════════════════════════════════════════════════════════════════
    # CAPÍTULO 4 — CONFIGURADOR
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph("4. Configurador de aplicación", style="Heading 1")

    doc.add_paragraph(
        "El configurador permite ajustar cada aspecto de la aplicación a través "
        "de 6 pestañas. Cada aplicación tiene su propia configuración independiente."
    )

    _figure(
        doc,
        "[SCREENSHOT: Configurador con las 6 pestañas visibles]",
        "Figura 8. Configurador de aplicación — vista general.",
    )

    # 4.1 General
    doc.add_paragraph("4.1 Pestaña General", style="Heading 2")
    doc.add_paragraph(
        "Configuración básica de la aplicación: nombre, descripción, color identificativo, "
        "opciones de auto-transferencia, detección de páginas en blanco y barcode manual."
    )

    _table(doc, ["Opción", "Descripción"], [
        ["Nombre", "Nombre de la aplicación que aparece en el Launcher"],
        ["Descripción", "Texto descriptivo del propósito de la aplicación"],
        ["Color", "Color identificativo para la aplicación en la lista"],
        ["Auto-transferencia", "Si está activo, transfiere automáticamente al completar el pipeline"],
        ["Detección de blancos", "Detecta y excluye automáticamente las páginas en blanco"],
        ["Umbral de blancos", "Porcentaje de píxeles blancos para considerar página en blanco"],
        ["Barcode manual", "Permite definir expresiones regulares o valores fijos para barcodes manuales"],
    ], "Tabla 3. Opciones de la pestaña General.")

    _figure(
        doc,
        "[SCREENSHOT: Pestaña General del configurador con todas las opciones]",
        "Figura 9. Pestaña General del configurador.",
    )

    # 4.2 Imagen
    doc.add_paragraph("4.2 Pestaña Imagen", style="Heading 2")
    doc.add_paragraph(
        "Define el formato de captura de las imágenes escaneadas: formato de fichero, "
        "modo de color (color, escala de grises, blanco y negro) y opciones de compresión."
    )

    _figure(
        doc,
        "[SCREENSHOT: Pestaña Imagen con selector de formato y compresión]",
        "Figura 10. Pestaña Imagen del configurador.",
    )

    # 4.3 Campos de lote
    doc.add_paragraph("4.3 Pestaña Campos de lote", style="Heading 2")
    doc.add_paragraph(
        "Permite definir campos dinámicos que el operador rellenará al crear o "
        "durante el procesamiento de un lote. Los tipos disponibles son: "
        "texto, fecha, lista desplegable y numérico."
    )

    _table(doc, ["Tipo de campo", "Widget en el Workbench", "Configuración"], [
        ["Texto", "Campo de texto libre", "Sin configuración adicional"],
        ["Fecha", "Selector de fecha con calendario", "Formato de fecha (dd/MM/yyyy, etc.)"],
        ["Lista", "Desplegable con opciones predefinidas", "Valores separados por coma"],
        ["Numérico", "Spinner numérico", "Mínimo, máximo y paso"],
    ], "Tabla 4. Tipos de campo de lote disponibles.")

    _figure(
        doc,
        "[SCREENSHOT: Pestaña Campos con varios campos definidos de distintos tipos]",
        "Figura 11. Pestaña Campos de lote con campos de ejemplo.",
    )

    # 4.4 Pipeline
    doc.add_paragraph("4.4 Pestaña Pipeline", style="Heading 2")
    doc.add_paragraph(
        "El pipeline es el corazón de DocScan Studio. Cada página pasa por una secuencia "
        "configurable de pasos que se ejecutan en orden. Se pueden añadir, eliminar, "
        "reordenar y activar/desactivar pasos individualmente."
    )

    _table(doc, ["Tipo de paso", "Descripción", "Ejemplo de uso"], [
        ["Imagen (image_op)", "Operación de procesamiento de imagen", "AutoDeskew, Crop, FxGrayscale"],
        ["Barcode", "Lectura de códigos de barras", "Detectar Code128 en zona superior"],
        ["OCR", "Reconocimiento óptico de caracteres", "Extraer texto con RapidOCR"],
        ["Script", "Lógica Python personalizada", "Separar documentos por barcode"],
    ], "Tabla 5. Tipos de paso del pipeline.")

    _figure(
        doc,
        "[SCREENSHOT: Pestaña Pipeline con varios pasos configurados y botones de acción]",
        "Figura 12. Editor de pipeline con pasos de ejemplo.",
    )

    doc.add_paragraph("4.4.1 Operaciones de imagen disponibles", style="Heading 3")
    _table(doc, ["Operación", "Descripción", "Parámetros principales"], [
        ["AutoDeskew", "Corrección automática de inclinación", "—"],
        ["ConvertTo1Bpp", "Conversión a blanco y negro (1 bit)", "threshold"],
        ["Crop", "Recorte rectangular", "x, y, w, h"],
        ["CropWhiteBorders", "Recorte automático de márgenes blancos", "tolerance"],
        ["CropBlackBorders", "Recorte automático de márgenes negros", "tolerance"],
        ["Resize", "Redimensionar", "width, height"],
        ["Rotate", "Rotar 90°, 180°, 270°", "angle"],
        ["RotateAngle", "Rotación libre por ángulo", "angle"],
        ["SetBrightness", "Ajustar brillo", "value (-100 a +100)"],
        ["SetContrast", "Ajustar contraste", "value (-100 a +100)"],
        ["RemoveLines", "Eliminar líneas horizontales/verticales", "direction"],
        ["FxDespeckle", "Eliminar ruido (speckle)", "—"],
        ["FxGrayscale", "Convertir a escala de grises", "—"],
        ["FxNegative", "Invertir colores", "—"],
        ["FxDilate", "Dilatación morfológica", "kernel_size"],
        ["FxErode", "Erosión morfológica", "kernel_size"],
        ["FxEqualizeIntensity", "Ecualización de histograma", "—"],
        ["FloodFill", "Relleno por inundación", "x, y, color"],
        ["RemoveHolePunch", "Eliminar agujeros de perforadora", "—"],
        ["SetResolution", "Ajustar DPI", "dpi"],
        ["SwapColor", "Intercambiar colores", "from_color, to_color"],
        ["KeepChannel", "Extraer canal de color", "channel (R/G/B)"],
        ["RemoveChannel", "Eliminar canal de color", "channel (R/G/B)"],
    ], "Tabla 6. Operaciones de imagen disponibles (23 operaciones).")

    doc.add_paragraph("4.4.2 Diálogo de paso Barcode", style="Heading 3")
    doc.add_paragraph(
        "El diálogo de configuración de barcode permite seleccionar el motor de detección, "
        "las simbologías a buscar, filtros regex, orientaciones y una ventana de búsqueda opcional."
    )

    _figure(
        doc,
        "[SCREENSHOT: Diálogo de configuración de paso Barcode]",
        "Figura 13. Configuración de paso Barcode.",
    )

    doc.add_paragraph("4.4.3 Diálogo de paso OCR", style="Heading 3")
    _figure(
        doc,
        "[SCREENSHOT: Diálogo de configuración de paso OCR]",
        "Figura 14. Configuración de paso OCR.",
    )

    doc.add_paragraph("4.4.4 Diálogo de paso Script", style="Heading 3")
    doc.add_paragraph(
        "El editor de scripts incluye plantilla por defecto, nombre de función configurable "
        "y posibilidad de editar en VS Code con autocompletado."
    )

    _figure(
        doc,
        "[SCREENSHOT: Diálogo de paso Script con código de ejemplo]",
        "Figura 15. Editor de paso Script con plantilla.",
    )

    doc.add_paragraph("4.4.5 Probar pipeline", style="Heading 3")
    doc.add_paragraph(
        "El botón «Probar pipeline» permite ejecutar el pipeline sobre una imagen "
        "de muestra y ver el resultado paso a paso, incluyendo la imagen resultante "
        "después de cada transformación."
    )

    _figure(
        doc,
        "[SCREENSHOT: Diálogo de resultados del test de pipeline]",
        "Figura 16. Resultados del test de pipeline paso a paso.",
    )

    # 4.5 Eventos
    doc.add_paragraph("4.5 Pestaña Eventos", style="Heading 2")
    doc.add_paragraph(
        "Los eventos son puntos de entrada del ciclo de vida donde se puede ejecutar "
        "código Python personalizado. Cada evento se dispara en un momento específico "
        "del flujo de trabajo."
    )

    _table(doc, ["Evento", "Cuándo se ejecuta"], [
        ["on_app_start", "Al abrir la aplicación en el workbench"],
        ["on_app_end", "Al cerrar la aplicación"],
        ["on_import", "Al pulsar Procesar (puede reemplazar la carga estándar)"],
        ["on_scan_complete", "Después de que el pipeline procese todas las páginas"],
        ["on_transfer_validate", "Antes de transferir; retornar False cancela"],
        ["on_transfer_advanced", "Transferencia avanzada completamente scripteada"],
        ["on_transfer_page", "Después de copiar cada página individual"],
        ["on_navigate_prev/next", "Al navegar entre páginas (personalizable)"],
        ["on_key_event", "Al pulsar una tecla personalizada"],
        ["init_global", "Al iniciar el programa (script global)"],
        ["verification_panel", "Define el panel de verificación personalizado"],
    ], "Tabla 7. Eventos del ciclo de vida.")

    _figure(
        doc,
        "[SCREENSHOT: Pestaña Eventos con código de on_scan_complete]",
        "Figura 17. Editor de eventos con código de ejemplo.",
    )

    # 4.6 Transferencia
    doc.add_paragraph("4.6 Pestaña Transferencia", style="Heading 2")
    doc.add_paragraph(
        "Configura cómo y dónde se exportan los documentos procesados. "
        "Soporta transferencia simple (a carpeta) y avanzada (mediante script)."
    )

    _table(doc, ["Opción", "Descripción"], [
        ["Carpeta destino", "Ruta donde se copiarán los ficheros"],
        ["Patrón de nombres", "Variables como {batch_id}, {page_index}, {fecha}, campos de lote"],
        ["Formato de salida", "TIFF, JPEG, PNG, BMP, PDF, PDF/A"],
        ["DPI de salida", "Resolución del fichero exportado"],
        ["Modo de color", "Color, escala de grises, blanco y negro"],
        ["Política de colisión", "Sufijo numérico, sobreescribir o fusionar (PDF/TIFF multi-página)"],
    ], "Tabla 8. Opciones de transferencia.")

    _figure(
        doc,
        "[SCREENSHOT: Pestaña Transferencia con todas las opciones configuradas]",
        "Figura 18. Pestaña Transferencia del configurador.",
    )

    # ════════════════════════════════════════════════════════════════
    # CAPÍTULO 5 — PIPELINE
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph("5. Pipeline en detalle", style="Heading 1")

    doc.add_paragraph(
        "Este capítulo profundiza en el funcionamiento interno del pipeline, "
        "el contexto disponible para scripts y las técnicas avanzadas de control de flujo."
    )

    doc.add_paragraph("5.1 Contexto del pipeline", style="Heading 2")
    doc.add_paragraph(
        "Cada script del pipeline recibe cuatro objetos de contexto que permiten "
        "acceder y modificar toda la información del procesamiento."
    )

    _table(doc, ["Objeto", "Tipo", "Descripción"], [
        ["app", "AppContext", "Nombre, descripción y configuración de la aplicación"],
        ["batch", "BatchContext", "ID, estado, campos y contadores del lote actual"],
        ["page", "PageContext", "Imagen, barcodes, OCR, campos, flags de la página"],
        ["pipeline", "PipelineContext", "Control de flujo: skip, abort, repeat, metadata"],
    ], "Tabla 9. Objetos de contexto disponibles en scripts.")

    doc.add_paragraph("5.2 Control de flujo", style="Heading 2")
    _table(doc, ["Método", "Efecto"], [
        ["pipeline.skip_step(id)", "Salta el paso indicado sin ejecutarlo"],
        ["pipeline.skip_to(id)", "Salta todos los pasos hasta llegar al indicado"],
        ["pipeline.abort(razón)", "Detiene el pipeline y marca la página para revisión"],
        ["pipeline.repeat_step(id)", "Vuelve a ejecutar un paso (máximo 3 veces por defecto)"],
        ["pipeline.replace_image(img)", "Reemplaza la imagen actual del pipeline"],
        ["pipeline.set_metadata(k, v)", "Almacena datos entre pasos"],
        ["pipeline.get_metadata(k)", "Recupera datos almacenados entre pasos"],
    ], "Tabla 10. Métodos de control de flujo del PipelineContext.")

    doc.add_paragraph("5.3 Ejemplo: separación por barcode", style="Heading 2")
    doc.add_paragraph(
        "Un caso de uso muy habitual es separar documentos automáticamente cuando "
        "se detecta una hoja separadora con un código de barras. El siguiente script "
        "asigna un ID de documento basado en el valor del barcode:"
    )
    doc.add_paragraph(
        'def separar(app, batch, page, pipeline):\n'
        '    seps = [b for b in page.barcodes if b.symbology == "CODE128"]\n'
        '    if seps:\n'
        '        page.fields["id_documento"] = seps[0].data\n'
        '        pipeline.set_metadata("ultimo_id", seps[0].data)\n'
        '    else:\n'
        '        page.fields["id_documento"] = pipeline.get_metadata("ultimo_id") or "SIN_ID"',
        style="Normal",
    )

    # ════════════════════════════════════════════════════════════════
    # CAPÍTULO 6 — SCRIPTING
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph("6. Scripting Python", style="Heading 1")

    doc.add_paragraph(
        "DocScan Studio incluye un motor de scripting Python completo que permite "
        "personalizar cualquier aspecto del procesamiento. Los scripts se compilan "
        "una sola vez al cargar la aplicación y se ejecutan con timeout configurable."
    )

    doc.add_paragraph("6.1 Variables disponibles", style="Heading 2")
    _table(doc, ["Variable", "Tipo", "Descripción"], [
        ["app", "AppContext", "Contexto de la aplicación actual"],
        ["batch", "BatchContext", "Contexto del lote actual"],
        ["page", "PageContext", "Contexto de la página (imagen, barcodes, OCR, campos, flags)"],
        ["pipeline", "PipelineContext", "Control de flujo (solo en ScriptStep, no en eventos)"],
        ["log", "Logger", "Logger para escribir mensajes al panel de log"],
        ["http", "httpx.Client", "Cliente HTTP para llamadas a APIs externas"],
        ["re", "módulo", "Expresiones regulares"],
        ["json", "módulo", "Serialización JSON"],
        ["datetime", "módulo", "Fecha y hora"],
        ["Path", "clase", "Manejo de rutas de fichero"],
    ], "Tabla 11. Variables disponibles en todos los scripts.")

    doc.add_paragraph("6.2 Buenas prácticas", style="Heading 2")
    _bullets(doc, [
        "Usar log.info(), log.warning() y log.error() en vez de print()",
        "Acceder a barcodes siempre via page.barcodes (nunca como objeto independiente)",
        "Los errores en scripts no detienen el pipeline: se registran en page.flags",
        "Usar pipeline.set_metadata() para compartir datos entre pasos",
        "El timeout por defecto es 30 segundos (configurable en settings)",
    ])

    doc.add_paragraph("6.3 Recetas comunes", style="Heading 2")

    doc.add_paragraph("6.3.1 Validar campos antes de transferir", style="Heading 3")
    doc.add_paragraph(
        '# Evento: on_transfer_validate\n'
        'def on_transfer_validate(app, batch, page):\n'
        '    if not batch.fields.get("numero_factura"):\n'
        '        raise ValueError("El número de factura es obligatorio")',
    )

    doc.add_paragraph("6.3.2 Llamar a una API externa", style="Heading 3")
    doc.add_paragraph(
        '# Paso Script en el pipeline\n'
        'def clasificar(app, batch, page, pipeline):\n'
        '    resp = http.post("https://api.ejemplo.com/clasificar",\n'
        '                     json={"text": page.ocr_text})\n'
        '    if resp.status_code == 200:\n'
        '        page.fields["tipo_doc"] = resp.json()["tipo"]',
    )

    doc.add_paragraph("6.3.3 Renombrar fichero de salida con OCR", style="Heading 3")
    doc.add_paragraph(
        '# Evento: on_transfer_page\n'
        'def on_transfer_page(app, batch, page):\n'
        '    # Extraer número de factura del texto OCR\n'
        '    match = re.search(r"FACTURA[:\\s]*(\\w+)", page.ocr_text)\n'
        '    if match:\n'
        '        page.fields["subdirectory"] = f"facturas/{match.group(1)}"',
    )

    # ════════════════════════════════════════════════════════════════
    # CAPÍTULO 7 — WORKBENCH
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph("7. Workbench", style="Heading 1")

    doc.add_paragraph(
        "El Workbench es la ventana de explotación donde se realiza el trabajo "
        "diario de captura y procesamiento. Se compone de cinco áreas principales."
    )

    _figure(
        doc,
        "[SCREENSHOT: Workbench completo con todas las áreas señaladas]",
        "Figura 19. Workbench de DocScan Studio — vista completa.",
    )

    doc.add_paragraph("7.1 Barra de herramientas", style="Heading 2")
    doc.add_paragraph(
        "La barra superior contiene los botones de escaneo, importación, transferencia, "
        "cierre de lote y opciones de tema. Incluye la casilla de configuración de escáner."
    )

    _figure(
        doc,
        "[SCREENSHOT: Barra de herramientas del Workbench con tooltips visibles]",
        "Figura 20. Barra de herramientas del Workbench.",
    )

    doc.add_paragraph("7.2 Panel de miniaturas", style="Heading 2")
    doc.add_paragraph(
        "El panel izquierdo muestra las miniaturas de todas las páginas del lote "
        "con bordes de color según su estado: procesada (verde), pendiente de revisión "
        "(naranja), excluida (gris), con error (rojo)."
    )

    _figure(
        doc,
        "[SCREENSHOT: Panel de miniaturas con páginas en distintos estados de color]",
        "Figura 21. Panel de miniaturas con estados visuales.",
    )

    doc.add_paragraph("7.3 Visor de documentos", style="Heading 2")
    doc.add_paragraph(
        "El visor central muestra la imagen de la página seleccionada con soporte "
        "para zoom, pan, rotación y overlays de barcodes y regiones OCR."
    )

    _figure(
        doc,
        "[SCREENSHOT: Visor con overlays de barcodes y regiones OCR marcadas]",
        "Figura 22. Visor con overlays de barcodes y OCR.",
    )

    doc.add_paragraph("7.4 Panel de barcodes", style="Heading 2")
    doc.add_paragraph(
        "El panel derecho superior muestra los barcodes detectados en la página actual "
        "con su valor, simbología, motor utilizado y rol asignado. Los puntos de color "
        "corresponden con los overlays del visor."
    )

    _figure(
        doc,
        "[SCREENSHOT: Panel de barcodes con varios códigos detectados]",
        "Figura 23. Panel de barcodes con códigos detectados.",
    )

    doc.add_paragraph("7.5 Panel de metadatos", style="Heading 2")
    doc.add_paragraph(
        "El panel inferior derecho tiene pestañas para los campos del lote, "
        "el panel de verificación personalizado (si está configurado) y el log "
        "en tiempo real del procesamiento."
    )

    _figure(
        doc,
        "[SCREENSHOT: Panel de metadatos con campos de lote rellenados]",
        "Figura 24. Panel de metadatos con campos de lote.",
    )

    doc.add_paragraph("7.6 Escaneo", style="Heading 2")
    doc.add_paragraph(
        "El botón de escaneo inicia la adquisición de imágenes desde el escáner "
        "configurado. Si se marca la casilla «Configurar escáner», se abre el "
        "diálogo de configuración SANE/TWAIN antes de escanear."
    )

    _figure(
        doc,
        "[SCREENSHOT: Diálogo de configuración del escáner SANE]",
        "Figura 25. Configuración del escáner.",
    )

    doc.add_paragraph("7.7 Importación de ficheros", style="Heading 2")
    doc.add_paragraph(
        "Se pueden importar ficheros mediante el botón «Importar» o arrastrando "
        "ficheros directamente sobre el visor. Formatos soportados: TIFF (incluido "
        "multi-página), JPEG, PNG, BMP y PDF."
    )

    # ════════════════════════════════════════════════════════════════
    # CAPÍTULO 8 — GESTIÓN DE LOTES
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph("8. Gestión de lotes", style="Heading 1")

    doc.add_paragraph(
        "El Gestor de Lotes permite consultar el histórico de todos los lotes "
        "procesados, filtrar por aplicación, fecha y estación, y reabrir lotes "
        "anteriores para revisión o reprocesamiento."
    )

    _figure(
        doc,
        "[SCREENSHOT: Gestor de lotes con filtros y lista de lotes]",
        "Figura 26. Gestor de Lotes con filtros aplicados.",
    )

    doc.add_paragraph("8.1 Estados del lote", style="Heading 2")
    _table(doc, ["Estado", "Descripción"], [
        ["Abierto", "Lote en proceso de captura o edición"],
        ["Procesado", "Pipeline completado, pendiente de transferencia"],
        ["Transferido", "Documentos exportados al destino"],
        ["Error", "Error durante el procesamiento"],
    ], "Tabla 12. Estados del ciclo de vida de un lote.")

    doc.add_paragraph("8.2 Reabrir y reprocesar", style="Heading 2")
    doc.add_paragraph(
        "Un lote cerrado se puede reabrir en el Workbench desde el Gestor de Lotes. "
        "Una vez abierto, se puede reprocesar cualquier página mediante Ctrl+P o "
        "reprocesar todas las páginas pendientes automáticamente."
    )

    # ════════════════════════════════════════════════════════════════
    # CAPÍTULO 9 — TRANSFERENCIA
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph("9. Transferencia", style="Heading 1")

    doc.add_paragraph("9.1 Transferencia simple (a carpeta)", style="Heading 2")
    doc.add_paragraph(
        "La transferencia simple copia los ficheros procesados a una carpeta de destino "
        "con un patrón de nombres configurable. Se pueden usar variables como "
        "{batch_id}, {page_index}, {fecha} y cualquier campo de lote."
    )

    doc.add_paragraph("9.2 Formatos de salida", style="Heading 2")
    _table(doc, ["Formato", "Extensión", "Características"], [
        ["TIFF", ".tif", "Compresión LZW/JPEG/Deflate, soporte multi-página"],
        ["JPEG", ".jpg", "Calidad configurable (1-100)"],
        ["PNG", ".png", "Sin pérdida, transparencia"],
        ["BMP", ".bmp", "Sin compresión"],
        ["PDF", ".pdf", "Documento PDF estándar"],
        ["PDF/A", ".pdf", "PDF/A-1b y PDF/A-2b para archivo a largo plazo"],
    ], "Tabla 13. Formatos de salida disponibles.")

    doc.add_paragraph("9.3 Política de colisión de nombres", style="Heading 2")
    _table(doc, ["Política", "Comportamiento"], [
        ["Sufijo", "Añade _001, _002... al nombre si ya existe"],
        ["Sobreescribir", "Reemplaza el fichero existente"],
        ["Fusionar", "Añade páginas al PDF/TIFF existente (multi-página)"],
    ], "Tabla 14. Políticas de colisión de nombres.")

    doc.add_paragraph("9.4 Transferencia avanzada (script)", style="Heading 2")
    doc.add_paragraph(
        "Para necesidades complejas, se puede escribir un script Python completo "
        "en el evento on_transfer_advanced que controla todo el proceso de exportación."
    )

    # ════════════════════════════════════════════════════════════════
    # CAPÍTULO 10 — AI MODE
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph("10. AI Mode", style="Heading 1")

    doc.add_paragraph(
        "AI Mode es un asistente conversacional integrado en el Launcher que permite "
        "crear y configurar aplicaciones mediante lenguaje natural. Utiliza Claude "
        "(Anthropic) o GPT-4o (OpenAI) según la API key configurada."
    )

    _figure(
        doc,
        "[SCREENSHOT: Panel AI Mode con una conversación de ejemplo]",
        "Figura 27. AI Mode — asistente conversacional.",
    )

    doc.add_paragraph("10.1 Configurar API key", style="Heading 2")
    doc.add_paragraph(
        "Para usar AI Mode, es necesario introducir una API key de Anthropic u OpenAI. "
        "Las claves se cifran con Fernet (AES-128-CBC) y se almacenan de forma segura "
        "en el directorio de configuración del usuario."
    )

    doc.add_paragraph("10.2 Capacidades", style="Heading 2")
    _bullets(doc, [
        "Crear nuevas aplicaciones describiendo el caso de uso",
        "Modificar la configuración de aplicaciones existentes",
        "Generar pipelines completos a partir de la descripción del documento",
        "Escribir scripts para lógica de negocio personalizada",
        "Documentar variables de transferencia y campos disponibles",
    ])

    doc.add_paragraph("10.3 Pipeline Assistant", style="Heading 2")
    doc.add_paragraph(
        "Además del AI Mode del Launcher, cada aplicación tiene un Pipeline Assistant "
        "accesible desde el configurador. Este asistente tiene contexto específico "
        "del pipeline configurado y puede sugerir mejoras o ayudar a depurar problemas."
    )

    _figure(
        doc,
        "[SCREENSHOT: Pipeline Assistant en el configurador]",
        "Figura 28. Pipeline Assistant integrado en el configurador.",
    )

    # ════════════════════════════════════════════════════════════════
    # CAPÍTULO 11 — ATAJOS DE TECLADO
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph("11. Referencia de atajos de teclado", style="Heading 1")

    _table(doc, ["Atajo", "Acción", "Contexto"], [
        ["Ctrl+S", "Escanear", "Workbench"],
        ["Ctrl+I", "Importar ficheros", "Workbench"],
        ["Ctrl+T", "Transferir lote", "Workbench"],
        ["Ctrl+P", "Reprocesar página", "Workbench"],
        ["← / →", "Página anterior / siguiente", "Workbench"],
        ["Home / End", "Primera / última página", "Workbench"],
        ["Ctrl+→", "Siguiente con barcode", "Workbench"],
        ["Ctrl+Shift+→", "Siguiente con revisión", "Workbench"],
        ["Ctrl++ / Ctrl+-", "Zoom in / out", "Visor"],
        ["Ctrl+F", "Ajustar a ventana", "Visor"],
        ["Ctrl+0", "Zoom 100%", "Visor"],
        ["R / Shift+R", "Rotar derecha / izquierda", "Visor"],
        ["M", "Marcar/desmarcar página", "Visor"],
        ["Del", "Eliminar página", "Visor"],
        ["Ctrl+W", "Cerrar lote", "Workbench"],
    ], "Tabla 15. Atajos de teclado del Workbench.")

    # ════════════════════════════════════════════════════════════════
    # CAPÍTULO 12 — GLOSARIO
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph("12. Glosario", style="Heading 1")

    _table(doc, ["Término", "Definición"], [
        ["Aplicación", "Perfil de procesamiento independiente con su pipeline, campos y transferencia"],
        ["Lote (Batch)", "Conjunto de páginas escaneadas/importadas en una sesión de trabajo"],
        ["Pipeline", "Secuencia ordenada de pasos que se ejecuta sobre cada página"],
        ["Paso (Step)", "Operación individual del pipeline: imagen, barcode, OCR o script"],
        ["Workbench", "Ventana de explotación donde se captura y procesa"],
        ["Launcher", "Ventana principal con la lista de aplicaciones"],
        ["Configurador", "Diálogo de configuración con 6 pestañas"],
        ["Overlay", "Marcadores visuales sobre el visor (barcodes, regiones OCR)"],
        ["Transferencia", "Proceso de exportar documentos procesados al destino"],
        ["AI Mode", "Asistente conversacional con IA para configurar aplicaciones"],
        ["Motor 1 (pyzbar)", "Motor de barcodes rápido con 12 simbologías"],
        ["Motor 2 (zxing-cpp)", "Motor de barcodes avanzado con 14 simbologías"],
        ["RapidOCR", "Motor OCR principal, offline, basado en ONNX"],
        ["SANE", "Sistema de acceso a escáneres en Linux"],
        ["TWAIN/WIA", "Interfaces de acceso a escáneres en Windows"],
        ["Fernet", "Algoritmo de cifrado para las API keys (AES-128-CBC + HMAC)"],
        ["WAL mode", "Write-Ahead Logging — modo de SQLite para concurrencia"],
    ], "Tabla 16. Glosario de términos.")

    # ── Guardar ──
    doc.save(OUTPUT)
    print(f"Manual generado: {OUTPUT}")
    print(f"Tamaño: {OUTPUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build()
